from docx import Document

from .base_parser import BaseParser


class DOCXParser(BaseParser):

    def parse(self, file_path: str) -> str:
        doc = Document(file_path)

        parts = []

        # Paragraphs
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)

        # Tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip()
                    for cell in row.cells
                    if cell.text.strip()
                )

                if row_text:
                    parts.append(row_text)

        return "\n".join(parts)